import json
import uuid

import docx
import requests
import math
import re
import time
import fitz


# import copy

# 定义函数,输入text,输出翻译:
def MStranslation_API(
        text,
        lang_in: str = None,
        lang_out="zh-Hans",
        subscription_key="pls input your MStranslation API key",
):
    """
    指定文本的微软翻译:
    args:
        text: 待翻译的原文本,或者列表;
        lang_in : 翻译前语言; 当None或者不包含时,自动辨识语言
        lang_out: 翻译后语言;可以列表形式,增加多种语言,例如:['ru','zh-Hans']
        subscription_key: 微软翻译API的key;
    out:
        trans_text: 即翻译后的文本字典列表,列表中每个字典包含了单个语种,或者多个语种的'text':'trans_text'的键值对;
        response: API接口输出,当翻译出错时,观察错误信息;
    """

    # Add your subscription key and endpoint
    subscription_key = subscription_key
    endpoint = "https://api.cognitive.microsofttranslator.com"  # 认知服务
    # Add your location, also known as region. The default is global.
    # This is required if using a Cognitive Services resource.
    location = "eastasia"
    path = "/translate"
    constructed_url = endpoint + path

    params = {"api-version": "3.0", "from": lang_in, "to": lang_out}

    headers = {
        "Ocp-Apim-Subscription-Key": subscription_key,
        "Ocp-Apim-Subscription-Region": location,
        "Content-type": "application/json",
        "X-ClientTraceId": str(uuid.uuid4()),
    }

    # You can pass more than one object in body.
    body = []

    if isinstance(text, list):
        for txt in text:
            body.append({"text": txt})

    if isinstance(text, str):
        body.append({"text": text})

    # print(body)
    # 当lang_out只有1种语言时,将lang_out变成列表:
    if isinstance(lang_out, str):
        lang_out = [lang_out]

    request = requests.post(constructed_url, params=params, headers=headers, json=body)
    response = request.json()
    # response_json = json.dumps(response, sort_keys=True, ensure_ascii=False, indent=4, separators=(',', ': '))
    # print(response_json)

    trans_text = []

    try:
        if isinstance(lang_out, list) and len(lang_out) > 1:
            trans_text = {}
            for i, lang in enumerate(lang_out):
                tmp = []
                for r in response:
                    tmp.append(r["translations"][i]["text"])
                trans_text[lang] = tmp
        else:
            for r in response:
                trans_text.append(r["translations"][0]["text"])
    except:
        print(response)

    return trans_text, response


def MStranslation_dynamicDictionary_API(
        text,
        dynamic_dict: dict = False,
        lang_in: str = None,
        lang_out: str = "zh-Hans",
        subscription_key="pls input your MStranslation API key",
):
    """
    指定文本的微软翻译:
    args:
        text: 待翻译的原文本,或者列表;
        dynamic_dict: 动态词典,包含有专有词汇,产品名称,人物人名等,已经有固定翻译的词汇,例:{'莫言':'Mr.Moyan'};缺省值:False,表示不需要动态词典;
        lang_in : 翻译前语言;当None或者不包含时,自动辨识语言
        lang_out: 翻译后语言;可以列表形式,增加多种语言,例如:['ru','zh-Hans']
        subscription_key: 微软翻译API的key;
    out:
        trans_text: 即翻译后的文本字典列表,列表中每个字典包含了单个语种,或者多个语种的'text':'trans_text'的键值对;
        response: API接口输出,当翻译出错时,观察错误信息;
    """

    # Add your subscription key and endpoint
    subscription_key = subscription_key
    endpoint = "https://api.cognitive.microsofttranslator.com"  # 认知服务
    # Add your location, also known as region. The default is global.
    # This is required if using a Cognitive Services resource.
    location = "eastasia"
    path = "/translate"
    constructed_url = endpoint + path

    params = {"api-version": "3.0", "from": lang_in, "to": lang_out}

    headers = {
        "Ocp-Apim-Subscription-Key": subscription_key,
        "Ocp-Apim-Subscription-Region": location,
        "Content-type": "application/json",
        "X-ClientTraceId": str(uuid.uuid4()),
    }

    # You can pass more than one object in body.
    body = []

    if isinstance(text, list):
        for txt in text:
            if isinstance(dynamic_dict, dict):
                for key in dynamic_dict.keys():
                    # 如果txt中含有动态词典的每个key的词汇,需将其替换成对应的value
                    sub_txt = (
                            '<mstrans:dictionary translation="'
                            + dynamic_dict[key]
                            + '">'
                            + key
                            + "</mstrans:dictionary>"
                    )
                    txt = re.sub(key, sub_txt, txt)
                body.append({"text": txt})
            elif dynamic_dict == False or dynamic_dict == "":
                body.append({"text": txt})
            else:
                print("Neither False nor Dictionary dynamic_dict is !")

    if isinstance(text, str):  # text为一条文本字符串
        if isinstance(dynamic_dict, dict):
            for key in dynamic_dict.keys():
                # 如果txt中含有动态词典的每个key的词汇,需将其替换成对应的value
                sub_txt = (
                        '<mstrans:dictionary translation="'
                        + dynamic_dict[key]
                        + '">'
                        + key
                        + "</mstrans:dictionary>"
                )
                text = re.sub(key, sub_txt, text)
            body.append({"text": text})
        elif dynamic_dict == False or dynamic_dict == "":
            body.append({"text": text})
        else:
            print("Neither False nor Dictionary dynamic_dict is !")

    # 当lang_out只有1种语言时,将lang_out变成列表:
    if isinstance(lang_out, str):
        lang_out = [lang_out]

    request = requests.post(constructed_url, params=params, headers=headers, json=body)
    response = request.json()
    # response_json = json.dumps(response, sort_keys=True, ensure_ascii=False, indent=4, separators=(',', ': '))
    # print(response_json)

    trans_text = []

    try:
        if isinstance(lang_out, list) and len(lang_out) > 1:
            trans_text = {}
            for i, lang in enumerate(lang_out):
                tmp = []
                for r in response:
                    tmp.append(r["translations"][i]["text"])
                trans_text[lang] = tmp
        else:
            for r in response:
                trans_text.append(r["translations"][0]["text"])
    except:
        print(response)

    return trans_text, response


# 定义一个函数,针对指定的paragraphs对象(存于内存中),对其每个run,予以判断某个run是否含有行间图形/图像, 行间图形图像,仅对其中的非空的text进行替代:
def paragraph_runs_replace(
        paragraph,
        text,
):
    """
    实现指定paragraph对象包含的每个run的仅text文字部分替代,图形,图像,空字符,不予处理;
    1. 指定paragraph对象,遍历每个run;获得run.text与paragraph.text之间的比例关系,并依据比例关系获得翻译后的段落text分布到每个run.text上的字符起止元组;
    2. 针对每个run,判断run是否是空字符;(当run.text是空字符时,可能包含图形,图像);当非空字符时,替换翻译text对应每个run的分布text
    注:该函数是调用python-docx库,该库是对document对象操作,是对object存储的指定内存进行操作,例如更改内容;当给对象赋值给新名称,实际上在内存中的位置不变化,因而,新名称与老名称指向内存同一位置,因而会操作在两个名称指向的同一个对象看到;因而,当使用深度复制的时候copy.deepcopy(),会出现新名称的内存,仍然存有未操作过的老的名称的对象;然而,paragraphs.text方法却能够看到操作过的text.
    所以,使用的局限性是:只能有一份内存位置的对象,以及针对对象的操作.不能够对对象复制.
    args:
        paragraph: python-docx的paragraph对象,可能来自于document,或table,或header,footer等;
        text: 将要替换的text,为单一字符串,非列表;
    out:
        paragraph: 输出完成操作的paragraph对象;
    """
    text_len = len(text)  # text是文本时,可以len()得字符数目;
    source_len = len(paragraph.text)

    run_attr = {}
    pointer = 0
    for i, r in enumerate(paragraph.runs):
        if source_len == 0:
            run_attr[i] = 0
        else:
            len_distrib = math.ceil(len(r.text) / source_len * text_len)  # 向上取整
            run_attr[i] = (pointer, pointer + len_distrib)
            pointer = pointer + len_distrib
    # print(run_attr)

    # 遍历run,判断非空字符;对非空字符的run,替换:
    if len(paragraph.runs) != 0:
        for i, r in enumerate(paragraph.runs):
            if r.text != "":
                # run文字的替换(run),格式保持不变:
                r.text = text[run_attr[i][0]: run_attr[i][1]]

    return paragraph


def Word_MStranslate(
        doc,
        dynamic_dict=False,
        lang_in='zh-Hans',
        lang_out='en',
        filename=False,
        subscription_key="pls input your MStranslation API key",
):
    """
    对Word文档中的段落,表格,页眉,页脚,予以翻译替换,格式不变,如有动态词典,文本中动态词典的key,直接替代成对应的value;
    A: 段落:
    1. 生成文档全部段落结构字典,以及text;
    2. 建立字典,对应text_dict的每个key,与其value在text列表中的index之间的一一对应关系;
    2. 调用翻译函数生成翻译后的trans_text列表;
    3. 对paragraphs的每一个段落,将翻译后的文本放回原文档;格式,样式不变;
    B: 表格翻译:
    1. 生成文档全部表结构的字典,以及text
    2. 建立字典,对应text_dict的每个key,与其value在text列表中的index之间的一一对应关系;
    3. 全部表结构的每个paragraph处理;
    C: 页眉,页脚,翻译替换
    args:
        doc: python-docx的document对象;
        lang_in : 翻译前语言;
        lang_out: 翻译后语言;一次翻译一种语言
        dynamic_dict: 动态词典,包含有专有词汇,产品名称,人物人名等,已经有固定翻译的词汇,例:{'莫言':'Mr.Moyan'}
        filename: 翻译后document对象,是否存入指定的文件路径;缺省:False时,不执行存盘动作;
        subscription_key: 微软翻译API的key;
    out:
        doc: 翻译后的,与原文样式/格式一致的document对象;
        当filename!=False时,指定路径位置存储翻译的文档;
        如翻译出错,则返回API输出的response;
    """
    # 文档段落部分翻译:
    # -----------------------------------
    # 生成待翻译的文本列表
    text_dict = {}
    for i, para in enumerate(doc.paragraphs):
        text_dict[(i)] = para.text
    # print(text_dict)
    text = list(text_dict.values())
    # 建立字典,对应text_dict的每个key,与其value在text列表中的index之间的一一对应关系;
    textindex_dict = {}
    for key in text_dict.keys():
        textindex_dict[key] = text.index(text_dict[key])

    # 生成翻译后的文本列表
    # 调用翻译API,判断是否带动态词典,并调用不同的API;
    if dynamic_dict == False:
        trans_text, _ = MStranslation_API(
            text, lang_in=lang_in, lang_out=lang_out, subscription_key=subscription_key
        )
    elif isinstance(dynamic_dict, dict):
        trans_text, _ = MStranslation_dynamicDictionary_API(
            text,
            dynamic_dict=dynamic_dict,
            lang_in=lang_in,
            lang_out=lang_out,
            subscription_key=subscription_key,
        )
    else:
        print("Neither False nor Dictionary dynamic_dict is !")

    # 对每一paragraph处理:
    for i, para in enumerate(doc.paragraphs):
        para_trans_text = trans_text[textindex_dict[(i)]]
        paragraph_runs_replace(para, para_trans_text)

    # -------------------------------------------

    # 表格部分翻译:
    # -------------------------------------
    # 获得全部表结构的字典,以及text
    text_dict = {}
    for t, table in enumerate(doc.tables):
        for r, row in enumerate(table.rows):
            for c, cell in enumerate(row.cells):
                for p, para in enumerate(cell.paragraphs):
                    text_dict[(t, r, c, p)] = para.text

    # print('text字典len:{}'.format(len(text)))
    text = list(text_dict.values())
    # 建立字典,对应text_dict的每个key,与其value在text列表中的index之间的一一对应关系;
    textindex_dict = {}
    for key in text_dict.keys():
        textindex_dict[key] = text.index(text_dict[key])

    # 调用翻译API,判断是否带动态词典,并调用不同的API;
    if dynamic_dict == False:
        trans_text, _ = MStranslation_API(
            text, lang_in=lang_in, lang_out=lang_out, subscription_key=subscription_key
        )
    elif isinstance(dynamic_dict, dict):
        trans_text, _ = MStranslation_dynamicDictionary_API(
            text,
            dynamic_dict=dynamic_dict,
            lang_in=lang_in,
            lang_out=lang_out,
            subscription_key=subscription_key,
        )
    else:
        print("Neither False nor Dictionary dynamic_dict is !")

    # 全部表结构的每个paragraph处理:
    for t, table in enumerate(doc.tables):
        for r, row in enumerate(table.rows):
            for c, cell in enumerate(row.cells):
                for p, para in enumerate(cell.paragraphs):
                    para_trans_text = trans_text[textindex_dict[(t, r, c, p)]]
                    paragraph_runs_replace(para, para_trans_text)

    # ----------------------------------
    # 页眉,页脚 翻译:
    # ----------------------------------
    # 获得全部section.header.paragraphs和secction.footer.paragraphs字典,以及text
    text_dict = {}
    for s, section in enumerate(doc.sections):
        for p, para in enumerate(section.header.paragraphs):  # 每个section只有一个header;
            text_dict[(s, "header", p)] = para.text
        for p, para in enumerate(section.footer.paragraphs):  # 每个section只有一个footer;
            if para.text.isdigit() == False and para.text != "":  # 页脚中动态页码,或者空白;不予处理;
                text_dict[(s, "footer", p)] = para.text
    # print(text_dict)
    text = list(text_dict.values())
    # 建立字典,对应text_dict的每个key,与其value在text列表中的index之间的一一对应关系;
    textindex_dict = {}
    for key in text_dict.keys():
        textindex_dict[key] = text.index(text_dict[key])

    # 调用翻译API,判断是否带动态词典,并调用不同的API;
    if dynamic_dict == False:
        trans_text, _ = MStranslation_API(
            text, lang_in=lang_in, lang_out=lang_out, subscription_key=subscription_key
        )
    elif isinstance(dynamic_dict, dict):
        trans_text, _ = MStranslation_dynamicDictionary_API(
            text,
            dynamic_dict=dynamic_dict,
            lang_in=lang_in,
            lang_out=lang_out,
            subscription_key=subscription_key,
        )
    else:
        print("Neither False nor Dictionary dynamic_dict is !")

    # 全部section.header.paragraphs的每个paragraph处理:
    for s, section in enumerate(doc.sections):
        for p, para in enumerate(section.header.paragraphs):  # 每个section只有一个header;
            para_trans_text = trans_text[textindex_dict[(s, "header", p)]]
            paragraph_runs_replace(para, para_trans_text)
        for p, para in enumerate(section.footer.paragraphs):  # 每个section只有一个header;
            if para.text.isdigit() == False and para.text != "":  # 页脚中动态页码,或者空白;不予处理;
                para_trans_text = trans_text[textindex_dict[(s, "footer", p)]]
                paragraph_runs_replace(para, para_trans_text)
    # ---------------------------------
    # 存盘,输出:
    if filename != False:
        doc.save(filename)
    return doc


# 定义一个函数实现,对源PDF文件,再生成一个PDF,其中text部分翻译,同时保持text颜色,大小,等style不变;image部分原封不动复制
# 采用get_text('dict')方法来处理image,优点是凡是页面上可以看到的图像,都会处理;缺点是对于重复image,会文件中插入;
# 检验是否含有中文字符
def isContainChinese(s):
    for c in s:
        if "\u4e00" <= c <= "\u9fa5":
            return True
    return False


# 检验是否全是中文字符


def isAllChinese(s):
    for c in s:
        if not ("\u4e00" <= c <= "\u9fa5"):
            return False
    return True


def PDF_MStranslate(
        input_path,
        output_path,
        lang_in="en",
        lang_out="zh-Hans",
        dynamic_dict=False,
        subscription_key="pls input your MS_translation API Key",
        image2txt=False,
        txtbox_borderColor="gray",
        out_font="cjk",
):
    """
    1. 读出每个image,包含inline image, 对每个image复制到新PDF;
    2. 读出每个block text,送入MS_translation API,获得trans_text列表;
    3. 获得每个block/line/span序列号,并读出对应的span_text的字体大小,颜色属性,字符长度;
    4. 将trans_text按照span_text对应的长度分布,分解每个span对应的span_trans_text分布;
    5. 将span_trans_text依照各自的起始位置,放入block/line/span, 字体颜色,大小不变;
    6. 将新生成的PDF存入out_path
    Args:
        input_path: 读入源PDF文件的路径;
        output_path:输出存盘的翻译PDF文件的路径;
        lang_in: 源PDF文本语言;
        lang_out: 翻译PDF文本语言;
        dynamic_dict: 动态词典,包含有专有词汇,产品名称,人物人名等,已经有固定翻译的词汇,例:{'莫言':'Mr.Moyan'};缺省值False,表示不需要动态词典;
        subscription_key: MS_translation API Key;
        image2txt: 是否在Image的方框中写入图像文件的txt,而不显示图像本身;缺省是False,即显示原图;
        txtbox_borderColor: 缺省给每个block画灰色的边框,即txtbox_borderColor='gray';当txtbox_borderColor=False时,不画边框;
        out_font: 翻译输出PDF显示的字符集; 中文设为:"cjk",英文可以设为:'helv';
    Out:
        操作PDF文档完成后,关闭打开的源PDF,新生成的PDF文档,存盘,结束.

    """

    t0 = time.time()
    # doc1 = fitz.open(sys.argv[1])
    doc1 = fitz.open(input_path)
    doc2 = fitz.open()
    pink = fitz.utils.getColor("pink")  # 给出颜色的RGB元组
    blue = fitz.utils.getColor("blue")
    # green = fitz.utils.getColor("green")
    if txtbox_borderColor != False:
        txtbox_borderColor = fitz.utils.getColor(txtbox_borderColor)
    else:
        txtbox_borderColor = None
    # gray = (0.9, 0.9, 0.9)

    # 增加一种font,显示多语言:
    font_cjk = fitz.Font(out_font)

    for page1 in doc1:
        # 生成同等大小的空白页
        page2 = doc2.new_page(
            -1, width=page1.mediabox_size[0], height=page1.mediabox_size[1]
        )
        # the text font we use
        # 也可以初始设置中文字体"china-s"
        # fontname=自定义的font名;字符集来自fontbuffer
        page2.insert_font(fontname="cjk", fontbuffer=font_cjk.buffer)

        # 绘矩形:
        img = page2.new_shape()  # prepare /Contents object
        # calculate /CropBox & displacement
        disp = fitz.Rect(
            page1.cropbox_position, page1.cropbox_position
        )  # 获得doc1的页面矩形原点坐标
        croprect = page1.rect + disp  # doc1的页面矩形相对原点位移;

        # draw original /CropBox rectangle
        img.draw_rect(croprect)
        img.finish(color=txtbox_borderColor, fill=None)

        # image 获取,写入pdf ,text 部分,extract,然后翻译,再放入PDF:
        # get_text('dict')方法,同get_text('blocks')方法相比,text部分丢失了换行符等控制字符,并且以span为单位,不能显示整个的block块
        blocks = page1.get_text("dict")
        blocks_ = page1.get_text("blocks")
        # 建立字典,获得各个span的text以及font,color
        span_attr = {}
        text = []
        for b, block in enumerate(blocks["blocks"]):
            if block["type"] == 1:  # block为image,缺点是无法获得smask:
                rect = fitz.Rect(block["bbox"])  # image的bbox坐标换成Rect对象
                if image2txt:
                    a = fitz.TEXT_ALIGN_CENTER
                    block_txt = blocks_[b][4]  # 该img的图像属性信息,即b[4]
                    rect += disp
                    img.draw_rect(rect)  # surround block rectangle
                    img.finish(width=0.3, color=pink)
                    img.insert_textbox(
                        rect, buffer=block_txt, fontsize=8, color=pink, align=a
                    )
                else:
                    page2.insert_image(
                        rect,
                        stream=block["image"],
                    )

            if block["type"] == 0:  # block为text
                block_txt = blocks_[b][4]
                # 需要将换行符给去了,否则,会导致翻译的txt也包含换行符,出现显示跨行
                block_txt = re.sub("\n", "", block_txt)
                block_txt_len = len(block_txt)  # block字符总长;
                text.append(block_txt)

                for l, line in enumerate(block["lines"]):
                    for s, span in enumerate(line["spans"]):
                        span_attr[(b, l, s, "size")] = span["size"]
                        span_attr[(b, l, s, "font")] = span["font"]
                        span_attr[(b, l, s, "color")] = span["color"]
                        span_attr[(b, l, s, "origin")] = span["origin"]
                        span_attr[(b, l, s, "text_index")] = len(text) - 1
                        span_attr[(b, l, s, "span_block_ratio")] = (
                                len(span["text"]) / block_txt_len
                        )

        try:
            trans_text, response = MStranslation_dynamicDictionary_API(
                text,
                dynamic_dict=dynamic_dict,
                lang_in=lang_in,
                lang_out=lang_out,
                subscription_key=subscription_key,
            )
        except:
            print(response)

        # 将trans_text按照各个span长/block长比例,分解至各个span, 获得span_attr[(b,l,s,'text')]

        for b, block in enumerate(blocks["blocks"]):
            if block["type"] == 0:  # block为text (image单独处理,因为这里无法获得smask)
                # block['bbox']为包含两个顶点的元组,Rect函数将元组变换为Rect对象
                rect = fitz.Rect(block["bbox"])
                # add dislacement of original /CropBox
                rect += disp

                img.draw_rect(rect)  # surround block rectangle
                a = fitz.TEXT_ALIGN_LEFT

                img.finish(width=0.3, color=txtbox_borderColor)

                pointer = 0
                for l, line in enumerate(block["lines"]):
                    for s, span in enumerate(line["spans"]):
                        # 将trans_text按照各个span长/block长比例,分解至各个span, 获得span_attr[(b,l,s,'text')]
                        span_transtxt_len = math.ceil(
                            len(trans_text[span_attr[(b, l, s, "text_index")]])
                            * span_attr[(b, l, s, "span_block_ratio")]
                        )
                        # print("span_transtxt_len:{}".format(span_transtxt_len))
                        # print(
                        #     "trans_text[{}]:{}".format(
                        #         b, trans_text[span_attr[(b, l, s, "text_index")]]
                        #     )
                        # )
                        span_attr[(b, l, s, "trans_text")] = trans_text[
                                                                 span_attr[(b, l, s, "text_index")]
                                                             ][pointer: pointer + span_transtxt_len]
                        pointer = pointer + span_transtxt_len
                        # print(
                        #     "span_trans_text:{}".format(span_attr[(b, l, s, "trans_text")])
                        # )

                        if rect.is_empty:  # do not rely on meaningful rects
                            print(
                                "skipping text of empty rect at ({}, {}) on page {}".format(
                                    rect.x0, rect.y0, page1.number
                                )
                            )
                        else:
                            # dict传出的color是sRGB格式,需要转换成RGB(R,G,B)元组格式(float(0,1)),用到这个内置函数:
                            color = fitz.utils.sRGB_to_pdf(
                                span_attr[(b, l, s, "color")]
                            )
                            point = fitz.Point(
                                span_attr[(b, l, s, "origin")]
                            )  # 将坐标元组,变换成坐标点;
                            # 根据字符串是否包含中文,来选择font
                            # if isContainChinese(span_attr[(b, l, s, "trans_text")]):
                            #     fontname = 'cjk'
                            # else:
                            #     fontname = "helv"

                            img.insert_text(
                                point=point,
                                buffer=span_attr[(b, l, s, "trans_text")],
                                # fontname="china-s",  # 中文font: 黑体
                                # fontname=span_attr[(b,l,s,'font')],
                                # fontname=fontname,
                                fontname="cjk",
                                fontsize=span_attr[(b, l, s, "size")],
                                color=color,
                            )  # 送入翻译后text

        img.commit()  # store /Contents of out page

    # save output file
    doc2.save(output_path, garbage=4, deflate=True, clean=True)
    doc1.close()
    doc2.close()  # 必要时,避免重复执行,重复操作;
    t1 = time.time()
    print("total time: {:.2f} sec".format((t1 - t0)))
